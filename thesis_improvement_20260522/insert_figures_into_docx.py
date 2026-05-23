from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


BASE = Path(__file__).resolve().parent
FIG = BASE / "figures"
SRC_DIR = Path.home() / "Desktop" / "毕业论文" / "第一版"
SRC = next(SRC_DIR.glob("*.docx"))
OUT = BASE / "毕业论文初版最终稿_增补版.docx"


def insert_after(paragraph, new_paragraph):
    paragraph._p.addnext(new_paragraph._p)
    return new_paragraph


def paragraph_after(doc, after, text="", style=None, alignment=None, font_size=None):
    p = doc.add_paragraph(style=style)
    if text:
        run = p.add_run(text)
        if font_size:
            run.font.size = Pt(font_size)
    if alignment is not None:
        p.alignment = alignment
    return insert_after(after, p)


def picture_after(doc, after, image_path, width=5.75):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    return insert_after(after, p)


def copy_style_from(source, target):
    target.style = source.style
    target.paragraph_format.alignment = source.paragraph_format.alignment
    target.paragraph_format.first_line_indent = source.paragraph_format.first_line_indent
    target.paragraph_format.left_indent = source.paragraph_format.left_indent
    target.paragraph_format.right_indent = source.paragraph_format.right_indent
    target.paragraph_format.space_before = source.paragraph_format.space_before
    target.paragraph_format.space_after = source.paragraph_format.space_after
    target.paragraph_format.line_spacing = source.paragraph_format.line_spacing


def body_after(doc, after, template, text):
    p = paragraph_after(doc, after, text=text)
    copy_style_from(template, p)
    return p


def caption_after(doc, after, template, text):
    p = paragraph_after(doc, after, text=text, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    copy_style_from(template, p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10.5)
    return p


def find_para(doc, exact_text, style_name=None):
    for p in doc.paragraphs:
        if p.text.strip() != exact_text:
            continue
        if style_name and p.style.name != style_name:
            continue
        return p
    raise ValueError(f"Paragraph not found: {exact_text!r}")


def clone_empty_paragraph_after(doc, after, template):
    p = doc.add_paragraph()
    copy_style_from(template, p)
    return insert_after(after, p)


def main():
    doc = Document(SRC)

    body_template = next(p for p in doc.paragraphs if p.style.name == "Body Text")
    caption_template = next(p for p in doc.paragraphs if p.text.strip().startswith("图4.1"))

    # Chapter 4: insert after the existing 4.2 pipeline explanation, before 4.3.
    anchor_4 = find_para(
        doc,
        "为了方便复现，项目提供了 `run_pipeline.py` 作为流水线入口。默认情况下，该脚本会依次执行数据模拟、特征工程、模型训练和实验评估。如果已有日志数据，也可以使用 `--skip-simulation` 参数跳过数据模拟，仅重建特征、模型和实验结果。",
    )
    p = body_after(
        doc,
        anchor_4,
        body_template,
        "在此基础上，实时监控模块进一步采用本地检测优先、智能解读异步补充的架构。该设计将风险检测链路与大模型自然语言生成链路分离，使风险分数、图表和可解释性结果能够优先返回，避免外部接口响应时间影响核心展示流程。",
    )
    p = picture_after(doc, p, FIG / "fig_detection_explanation_loop.png")
    p = caption_after(doc, p, caption_template, "图4.3 用户行为风险检测、解释与处置闭环")
    p = body_after(
        doc,
        p,
        body_template,
        "具体到在线监控流程，前端首先通过实时日志接口获取本地模型和规则融合后的检测结果，并立即渲染风险仪表盘、行为雷达图和 SHAP 解释图；若启用大模型模块，后端再创建异步解读任务，由前端根据任务编号增量刷新智能摘要和研判建议。",
    )
    p = picture_after(doc, p, FIG / "fig_async_llm_architecture.png")
    caption_after(doc, p, caption_template, "图4.4 实时监控模块的本地优先与异步智能解读架构")

    # Chapter 5: insert after backend startup check paragraph, before 5.5.
    anchor_5 = find_para(
        doc,
        "为了提高系统演示稳定性，后端增加启动检查模块。启动时会检查特征数据、模型文件、特征列表、实验结果、本地 ECharts 和 vis-network 文件以及 LLM 配置状态。若文件缺失，系统会给出清晰提示，便于快速定位问题。",
    )
    p = body_after(
        doc,
        anchor_5,
        body_template,
        "除单条日志检测外，后端还会从用户维度聚合历史风险信息，形成用户风险画像。画像结果综合异常次数、平均风险、最高风险、敏感访问和规则触发原因等信息，为用户详情页和后续人工复核提供连续证据。",
    )
    p = picture_after(doc, p, FIG / "fig_user_profile_aggregation.png")
    caption_after(doc, p, caption_template, "图5.1 用户风险画像聚合流程")

    # Chapter 6: place after the existing feature-importance analysis to avoid renumbering old Figure 6.1-6.4 captions.
    anchor_6 = find_para(
        doc,
        "特征重要性分析也说明本文特征工程设计较为有效。系统没有只依赖单一字段，而是从角色、位置、时间、操作、敏感资源和用户基线等多个角度综合判断风险。",
    )
    p = body_after(
        doc,
        anchor_6,
        body_template,
        "综合实验指标设计和模型结果可以看出，安全场景下的异常检测不能只关注准确率，还需要同时考虑漏报、误报、高危场景召回和实时推理成本。本文将通用分类指标、不平衡检测指标、安全业务指标和运行指标结合起来，用于评价模型在真实安全运营中的综合适用性。",
    )
    p = picture_after(doc, p, FIG / "fig_security_metric_framework.png")
    caption_after(doc, p, caption_template, "图6.5 面向安全业务的模型评价指标框架")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
