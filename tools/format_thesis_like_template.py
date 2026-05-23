from __future__ import annotations

import re
from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parent.parent
MD_PATH = ROOT / "毕业论文初稿.md"
DOCX_PATH = ROOT / "毕业论文初稿.docx"
FORMAL_DOCX_PATH = ROOT / "毕业论文初稿_正式版式.docx"
ASSET_DIR = ROOT / "论文初稿_assets"

CHINESE_FONT = "宋体"
HEADING_FONT = "黑体"
ENGLISH_FONT = "Times New Roman"


def set_rfonts(element, east_asia: str, latin: str | None = None) -> None:
    r_pr = element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    latin = latin or east_asia
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)


def set_run_font(run, east_asia: str = CHINESE_FONT, size: float = 12, bold: bool = False, latin: str | None = None) -> None:
    run.font.name = latin or east_asia
    set_rfonts(run._element, east_asia, latin)
    run.font.size = Pt(size)
    run.bold = bold


def set_style_font(style, east_asia: str, size: float, bold: bool = False, latin: str | None = None) -> None:
    style.font.name = latin or east_asia
    style.font.size = Pt(size)
    style.font.bold = bold
    set_rfonts(style._element, east_asia, latin)


def clear_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        paragraph._p.remove(run._element)


def set_page_number_start(section, start: int = 1, fmt: str | None = None) -> None:
    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), str(start))
    if fmt:
        pg_num.set(qn("w:fmt"), fmt)


def add_page_number(section, roman: bool = False) -> None:
    section.footer.is_linked_to_previous = False
    p = section.footer.paragraphs[0]
    clear_runs(p)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE \\* roman" if roman else "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_text = OxmlElement("w:t")
    fld_text.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_text)
    run._r.append(fld_end)
    set_run_font(run, CHINESE_FONT, 10.5)


def apply_page_setup(section) -> None:
    section.page_width = Pt(595.3)
    section.page_height = Pt(841.9)
    section.top_margin = Pt(70.9)
    section.bottom_margin = Pt(56.7)
    section.left_margin = Pt(85.0)
    section.right_margin = Pt(45.4)
    section.header_distance = Pt(42.5)
    section.footer_distance = Pt(42.5)


def enable_update_fields(doc: Document) -> None:
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def init_styles(doc: Document) -> None:
    for section in doc.sections:
        apply_page_setup(section)

    normal = doc.styles["Normal"]
    set_style_font(normal, CHINESE_FONT, 12, latin=ENGLISH_FONT)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    body = doc.styles["Body Text"] if "Body Text" in doc.styles else doc.styles.add_style("Body Text", 1)
    set_style_font(body, CHINESE_FONT, 12, latin=ENGLISH_FONT)
    body.paragraph_format.first_line_indent = Pt(24)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_before = Pt(0)
    body.paragraph_format.space_after = Pt(0)

    list_style = doc.styles["List Paragraph"]
    set_style_font(list_style, CHINESE_FONT, 12, latin=ENGLISH_FONT)
    list_style.paragraph_format.first_line_indent = Pt(24)
    list_style.paragraph_format.line_spacing = 1.5
    list_style.paragraph_format.space_before = Pt(0)
    list_style.paragraph_format.space_after = Pt(0)

    for style_name, size, before, after in [
        ("Heading 1", 15, 12, 12),
        ("Heading 2", 14, 8, 6),
        ("Heading 3", 12, 6, 4),
    ]:
        style = doc.styles[style_name]
        set_style_font(style, HEADING_FONT, size, True)
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_center(doc: Document, text: str, font: str, size: float, bold: bool = False, after: float = 8) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_run_font(r, font, size, bold, ENGLISH_FONT if font == ENGLISH_FONT else None)


def add_body(doc: Document, text: str, english: bool = False) -> None:
    p = doc.add_paragraph(style="Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    if english:
        set_run_font(r, ENGLISH_FONT, 12, False, ENGLISH_FONT)
    else:
        set_run_font(r, CHINESE_FONT, 12, False, ENGLISH_FONT)


def add_no_indent(doc: Document, text: str, size: float = 12, bold: bool = False, english: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    if english:
        set_run_font(r, ENGLISH_FONT, size, bold, ENGLISH_FONT)
    else:
        set_run_font(r, CHINESE_FONT, size, bold, ENGLISH_FONT)


def add_label_paragraph(doc: Document, label: str, rest: str, english: bool = False) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(0)
    r1 = p.add_run(label)
    r2 = p.add_run(rest)
    if english:
        set_run_font(r1, ENGLISH_FONT, 12, True, ENGLISH_FONT)
        set_run_font(r2, ENGLISH_FONT, 12, False, ENGLISH_FONT)
    else:
        set_run_font(r1, CHINESE_FONT, 12, True, ENGLISH_FONT)
        set_run_font(r2, CHINESE_FONT, 12, False, ENGLISH_FONT)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.line_spacing = 1.5
    size = 15 if level == 1 else 14 if level == 2 else 12
    for run in p.runs:
        set_run_font(run, HEADING_FONT, size, True)


def add_toc_field(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    placeholder = OxmlElement("w:t")
    placeholder.text = "打开 Word 后请更新目录域"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(placeholder)
    run._r.append(fld_end)
    set_run_font(run, CHINESE_FONT, 12)


def add_static_toc(doc: Document, headings: list[tuple[int, str]]) -> None:
    for level, text in headings:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.left_indent = Pt(0 if level == 1 else 24 if level == 2 else 48)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(text)
        set_run_font(r, CHINESE_FONT, 12 if level == 1 else 11)


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    if not headers:
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, CHINESE_FONT, 10.5, True, ENGLISH_FONT)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row[: len(headers)]):
            cell = cells[i]
            cell.text = ""
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.2
            r = p.add_run(text)
            set_run_font(r, CHINESE_FONT, 10, False, ENGLISH_FONT)
    doc.add_paragraph()


def add_image(doc: Document, image_name: str, caption: str) -> None:
    path = ASSET_DIR / image_name
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(6)
    p.add_run().add_picture(str(path), width=Inches(5.65))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Pt(0)
    cap.paragraph_format.space_after = Pt(6)
    r = cap.add_run(caption)
    set_run_font(r, CHINESE_FONT, 10.5, False, ENGLISH_FONT)


def extract_section(md: str, start_heading: str, next_heading_pattern: str) -> str:
    start = md.index(start_heading) + len(start_heading)
    match = re.search(next_heading_pattern, md[start:])
    end = start + match.start() if match else len(md)
    return md[start:end].strip()


def clean_lines(text: str) -> list[str]:
    return [line.strip() for line in text.splitlines() if line.strip()]


def collect_headings(md: str) -> list[tuple[int, str]]:
    body_start = md.index("# 1 绪论")
    headings: list[tuple[int, str]] = []
    for raw in md[body_start:].splitlines():
        line = raw.strip()
        if line.startswith("### "):
            headings.append((3, line[4:].strip()))
        elif line.startswith("## "):
            headings.append((2, line[3:].strip()))
        elif line.startswith("# "):
            headings.append((1, line[2:].strip()))
    return headings


def add_cover(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph()
    add_center(doc, "淮 安 大 学", HEADING_FONT, 26, True, 14)
    add_center(doc, "毕业设计（论文）", HEADING_FONT, 22, True, 34)
    add_center(doc, "学生姓名：__________    学号：__________", CHINESE_FONT, 12, False, 12)
    add_center(doc, "设计（论文）题目：基于用户行为数据评估的", CHINESE_FONT, 14, False, 4)
    add_center(doc, "企业信息数据安全防护系统", HEADING_FONT, 16, True, 28)
    add_center(doc, "专业：计算机相关专业    班级：__________", CHINESE_FONT, 12, False, 12)
    add_center(doc, "校内指导老师：__________", CHINESE_FONT, 12, False, 6)
    add_center(doc, "（姓    名）        （专业技术职务）", CHINESE_FONT, 12, False, 12)
    add_center(doc, "校外指导老师：    /        /", CHINESE_FONT, 12, False, 6)
    add_center(doc, "（姓    名）        （专业技术职务）", CHINESE_FONT, 12, False, 34)
    add_center(doc, "2026    年    05    月", CHINESE_FONT, 12, False, 0)


def add_front_matter(doc: Document, md: str) -> None:
    zh_abs = extract_section(md, "## 摘要", r"\n## Abstract")
    en_abs = extract_section(md, "## Abstract", r"\n# 1 ")
    zh_lines = clean_lines(zh_abs)
    zh_keywords = ""
    zh_paragraphs: list[str] = []
    for line in zh_lines:
        if line.startswith("关键词"):
            zh_keywords = line
        else:
            zh_paragraphs.append(line)
    add_center(doc, "摘要", HEADING_FONT, 16, True, 12)
    for paragraph in zh_paragraphs:
        add_body(doc, paragraph)
    if zh_keywords:
        label, rest = zh_keywords.split("：", 1) if "：" in zh_keywords else ("关键词", zh_keywords.replace("关键词", "", 1))
        add_label_paragraph(doc, f"{label}：", rest)
    doc.add_page_break()

    en_lines = clean_lines(en_abs)
    en_keywords = ""
    en_paragraphs: list[str] = []
    for line in en_lines:
        if line.startswith("Key words") or line.startswith("Keywords"):
            en_keywords = line
        else:
            en_paragraphs.append(line)
    add_center(doc, "Abstract", ENGLISH_FONT, 16, True, 12)
    for paragraph in en_paragraphs:
        add_body(doc, paragraph, english=True)
    if en_keywords:
        if ":" in en_keywords:
            label, rest = en_keywords.split(":", 1)
            add_label_paragraph(doc, f"{label}:", rest, english=True)
        else:
            add_no_indent(doc, en_keywords, bold=True, english=True)


def add_toc(doc: Document, headings: list[tuple[int, str]]) -> None:
    add_center(doc, "目录", HEADING_FONT, 16, True, 12)
    add_static_toc(doc, [(1, "摘要"), (1, "Abstract"), *headings])


def add_body_from_markdown(doc: Document, md: str) -> None:
    body_start = md.index("# 1 绪论")
    table_lines: list[str] = []
    in_code = False
    first_chapter = True

    def flush_table() -> None:
        nonlocal table_lines
        if not table_lines:
            return
        rows = []
        for line in table_lines:
            if re.match(r"^\|\s*-", line):
                continue
            rows.append([part.strip() for part in line.strip("|").split("|")])
        if rows:
            add_table(doc, rows[0], rows[1:])
        table_lines = []

    for raw in md[body_start:].splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("```"):
            in_code = not in_code
            continue
        if line.startswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        if line.startswith("# "):
            text = line[2:].strip()
            if first_chapter:
                first_chapter = False
            else:
                doc.add_page_break()
            add_heading(doc, text, 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("!["):
            match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if match:
                add_image(doc, Path(match.group(2)).name, match.group(1))
        elif re.match(r"^\[\d+\]", line):
            add_no_indent(doc, line, size=10.5)
        elif in_code:
            add_no_indent(doc, line, size=10.5)
        elif re.match(r"^表\d+(\.\d+)?\s+", line):
            add_center(doc, line, CHINESE_FONT, 10.5, False, 4)
        else:
            add_body(doc, line)
    flush_table()


def build_docx() -> None:
    md = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    init_styles(doc)
    enable_update_fields(doc)

    cover_section = doc.sections[0]
    apply_page_setup(cover_section)
    cover_section.footer.is_linked_to_previous = False
    clear_runs(cover_section.footer.paragraphs[0])
    add_cover(doc)

    front_section = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_page_setup(front_section)
    set_page_number_start(front_section, 1, "roman")
    add_page_number(front_section, roman=True)
    add_front_matter(doc, md)

    toc_section = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_page_setup(toc_section)
    set_page_number_start(toc_section, 1, "roman")
    add_page_number(toc_section, roman=True)
    add_toc(doc, collect_headings(md))

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    apply_page_setup(body_section)
    set_page_number_start(body_section, 1)
    add_page_number(body_section)
    add_body_from_markdown(doc, md)

    doc.save(FORMAL_DOCX_PATH)
    try:
        doc.save(DOCX_PATH)
        saved_main = True
    except PermissionError:
        saved_main = False

    with ZipFile(FORMAL_DOCX_PATH) as zf:
        media_count = len([n for n in zf.namelist() if n.startswith("word/media/")])
    print(f"written: {FORMAL_DOCX_PATH}")
    print(f"updated_main: {saved_main}")
    print(f"media={media_count}, paragraphs={len(doc.paragraphs)}, tables={len(doc.tables)}, sections={len(doc.sections)}")


if __name__ == "__main__":
    build_docx()
