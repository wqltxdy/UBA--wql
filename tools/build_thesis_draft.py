from __future__ import annotations

import json
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Cm, Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent.parent
ASSET_DIR = ROOT / "论文初稿_assets"
OUT_MD = ROOT / "毕业论文初稿.md"
OUT_DOCX = ROOT / "毕业论文初稿.docx"
ALT_DOCX = ROOT / "毕业论文初稿_样例格式.docx"


def load_support() -> dict:
    with open(ASSET_DIR / "support.json", "r", encoding="utf-8") as f:
        return json.load(f)


def pct(x: float) -> str:
    return f"{x * 100:.2f}%"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "宋体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    r.font.size = Pt(10)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        set_cell_shading(table.rows[0].cells[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)
    doc.add_paragraph()


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(1.6)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.space_after = Pt(0)

    for style_name, size, bold in [
        ("Title", 22, True),
        ("Heading 1", 15, True),
        ("Heading 2", 14, True),
        ("Heading 3", 12, True),
    ]:
        style = styles[style_name]
        style.font.name = "黑体" if bold else "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if bold else "宋体")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.paragraph_format.first_line_indent = Pt(0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_before = Pt(6)
        style.paragraph_format.space_after = Pt(6)

    if "Body Text" in styles:
        body = styles["Body Text"]
    else:
        body = styles.add_style("Body Text", 1)
    body.font.name = "宋体"
    body._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    body.font.size = Pt(12)
    body.paragraph_format.first_line_indent = Pt(24)
    body.paragraph_format.line_spacing = 1.5
    body.paragraph_format.space_after = Pt(0)


def add_para(doc: Document, text: str = "", style: str | None = None) -> None:
    p = doc.add_paragraph(style=style or "Body Text")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if text else WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(12)


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT if level > 1 else WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)


def add_image(doc: Document, name: str, caption: str, width: float = 5.8) -> None:
    path = ASSET_DIR / name
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cap.runs:
            run.font.name = "宋体"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            run.font.size = Pt(10.5)


def md_image(name: str, caption: str) -> str:
    return f"\n![{caption}](论文初稿_assets/{name})\n"


def build_markdown(support: dict) -> str:
    rows = support["experiment"]["model_compare"]
    lgb = next(r for r in rows if r["Model"] == "LightGBM")
    hybrid = next(r for r in rows if r["Model"] == "LightGBM + Expert Rules")
    lr = next(r for r in rows if r["Model"] == "Logistic Regression")
    iso = next(r for r in rows if r["Model"] == "Isolation Forest")
    text = f"""# 基于用户行为数据评估的企业信息数据安全防护系统

毕业设计（论文）初稿

题目：基于用户行为数据评估的企业信息数据安全防护系统

专业：计算机相关专业

学生姓名：__________

指导教师：__________

完成日期：2026年5月

## 摘要

随着企业办公系统、数据平台和业务应用逐步走向数字化、网络化和集中化，企业内部用户在登录认证、文件访问、敏感数据下载、远程连接和业务操作过程中会持续产生大量行为日志。这些日志蕴含着用户日常行为基线、业务角色权限边界以及潜在安全风险线索。传统安全防护体系往往更关注外部攻击、边界防护和规则告警，对内部用户行为偏离、账号异常使用、越权访问和敏感数据异常操作的识别能力不足。当企业数据资产不断增长时，仅依靠人工审计和静态规则难以及时发现高风险行为，也难以在告警产生后向安全人员提供充分解释。因此，构建一个面向企业信息数据安全防护的用户行为分析系统，具有较强的实际意义。

本文围绕“基于用户行为数据评估的企业信息数据安全防护系统”展开研究与实现，设计并开发了一套用户行为分析（User Behavior Analytics，UBA）系统。系统以模拟企业用户行为日志为基础，构建包含用户角色、登录位置、操作类型、访问时间、文件敏感级别、短时间操作频次、登录失败次数、角色操作匹配关系等多维特征的数据集；在算法层面采用 LightGBM 监督学习模型识别异常行为，同时引入面向安全场景的专家规则，对深夜高敏操作、周末高频操作、远程或未知位置叠加登录失败、角色操作不匹配叠加高敏资源访问等场景进行风险上调；在解释层面结合 SHAP 方法输出模型主要贡献特征，并通过规则命中说明补充业务安全知识；在展示层面基于 Flask、ECharts 和 vis-network 实现实时监控、异常用户表、用户画像、风险关系图和实验对比页面；在智能化增强层面，系统支持通过 DeepSeek / OpenAI 兼容接口生成自然语言风险解读、告警研判和用户画像总结。

实验结果表明，在时间切分的测试集上，LightGBM 模型取得了较好的通用机器学习指标，其 F1 值为 {lgb['F1-Score']:.4f}，PR-AUC 为 {lgb['PR-AUC']:.4f}；引入专家规则后的 LightGBM + Expert Rules 模型在通用指标上并非全面领先，但高危场景召回率由 {lgb['HighRisk-Recall']:.4f} 提升至 {hybrid['HighRisk-Recall']:.4f}，加权安全代价由 {lgb['Security-Cost']:.0f} 降低至 {hybrid['Security-Cost']:.0f}。这说明专家规则融合模型的价值并不在于单纯追求所有统计指标最优，而在于降低关键风险漏报、增强业务策略一致性和提高系统可解释性。本文最终形成了“检测、解释、研判、处置建议”的安全分析闭环，可为企业内部数据安全防护和用户行为风险审计提供参考。

关键词：用户行为分析；异常检测；LightGBM；专家规则；SHAP；数据安全；告警研判

## Abstract

With the continuous digitalization of enterprise information systems, user behavior logs generated by authentication, file access, sensitive data operations and remote connections have become important evidence for security risk assessment. Traditional security mechanisms mainly rely on boundary protection and static rules, which are often insufficient for identifying internal behavior deviations, privilege misuse and suspicious access to sensitive data. This thesis designs and implements a User Behavior Analytics system for enterprise information and data security protection. The system builds multi-dimensional behavioral features from simulated enterprise logs, trains a LightGBM model for anomaly detection, integrates expert rules for high-risk security scenarios, uses SHAP to explain model outputs, and provides a Flask-based visualization dashboard. In addition, the system supports DeepSeek or OpenAI-compatible large language model interfaces to generate natural-language alert interpretation, triage suggestions and user profile summaries.

Experimental results show that the LightGBM model performs well on general machine learning metrics. The hybrid model with expert rules improves high-risk recall and reduces weighted security cost, although some general metrics slightly decrease. The result indicates that expert rules are valuable for security-oriented risk control, especially for reducing missed detections in high-risk scenarios and improving interpretability. The system forms a closed loop from detection to explanation, triage and disposal recommendation, which provides a practical reference for enterprise data security protection.

Key words: User Behavior Analytics; Anomaly Detection; LightGBM; Expert Rules; SHAP; Data Security

# 1 绪论

## 1.1 研究背景

企业信息化建设不断深入，业务流程、办公协同、数据存储和权限管理越来越依赖统一的信息系统。用户在系统中的登录、查询、下载、修改、导出和远程访问等操作都会被记录为日志。与传统网络流量或主机告警相比，用户行为日志更贴近企业业务活动，能够反映用户身份、角色职责、访问时间、操作对象、资源敏感程度和行为频率等信息。因此，用户行为日志不仅是审计追踪的基础，也是识别内部风险、异常账号使用和数据泄露隐患的重要数据来源。Almohaimeed 等在供应链网络风险评估研究中指出，第三方依赖、外部协作和组织间安全能力差异会形成持续存在的“安全间隙”（security gaps）[1]。这一观点对企业内部数据安全同样具有启发意义：当业务系统、外部接口、外包人员、远程办公和跨部门数据流转共同存在时，风险往往不只来自单点漏洞，而是来自身份、权限、行为监测和处置流程之间的间隙。用户行为分析的价值正在于通过持续监测合法账号下的行为偏离，弥补静态权限控制和边界防护难以覆盖的安全间隙。

在企业安全防护场景中，外部攻击固然重要，但内部用户风险同样不可忽视。一方面，合法账号可能被盗用，攻击者利用正常账号绕过边界防护后进行横向移动和数据窃取；另一方面，内部人员也可能由于越权访问、违规下载或恶意泄露造成数据安全事件。相比明显的攻击流量，这类行为往往隐藏在正常业务访问之中，表现为行为模式的细微偏离，例如非工作时间访问高敏文件、短时间内大量下载、从陌生位置登录、登录失败次数异常增加、普通员工执行与角色不匹配的管理操作等。若系统无法对这些行为进行持续分析，就容易形成安全监控盲区。

传统安全检测系统通常依赖固定规则。固定规则具有实现简单、解释直观的优势，但面对复杂业务和不断变化的用户行为时，容易出现两个问题：第一，规则过严会带来大量误报，增加安全人员复核压力；第二，规则过松又会导致高风险行为漏报。机器学习方法能够从历史日志中学习行为统计规律，为异常检测提供更灵活的判断能力。然而，单纯依靠模型也存在解释不足和业务目标不完全一致的问题。例如，一个模型在整体 F1 或 AUC 指标上表现较好，并不必然意味着它对高危数据泄露场景具有更好的兜底能力。安全场景往往更关心高危漏报，而不是所有样本上的平均表现。

基于上述背景，本文设计并实现一套面向企业信息数据安全防护的 UBA 系统。系统并不简单停留在异常检测模型训练，而是进一步将专家规则、模型解释、可视化展示和智能化文本研判纳入整体设计，使系统能够从“是否异常”扩展到“为什么异常、属于什么场景、应如何处置”。这种闭环式设计更符合企业安全运营工作的实际需求。

## 1.2 研究意义

本文研究具有以下几个方面的意义。首先，从数据安全角度看，用户行为分析能够补充传统边界安全防护的不足。企业内部数据往往分布在文件系统、业务系统和数据库中，访问主体具有合法账号和业务权限，单纯依靠网络边界或病毒查杀难以及时发现风险。通过对用户行为日志进行建模，可以在用户访问敏感资源时识别其行为是否偏离日常模式，从而为内部风险发现提供支持。尤其在第三方接入、远程访问和多系统协作环境下，安全间隙并不一定表现为明显漏洞，而可能表现为权限边界模糊、账号行为缺少持续校验、异常访问缺少解释和处置建议等问题。本文系统通过行为建模、专家规则和告警研判机制，对这些间隙进行监测和解释。

其次，从安全运营角度看，系统需要兼顾检测效果和可解释性。安全人员面对告警时，不仅需要知道风险分数，还需要知道风险来自哪些证据。本文将 LightGBM 模型输出、专家规则命中结果和 SHAP 特征贡献结合起来，使告警解释既包含统计学习依据，也包含业务规则依据。相比单一分数，这种解释方式更便于安全人员复核和处置。

再次，从实验评价角度看，本文不仅使用 Accuracy、Precision、Recall、F1、AUC、PR-AUC 等通用指标，还设计了 HighRisk-Recall、Security-Cost 和 Security-Score 等面向安全业务的补充指标。由于数据安全场景中高危漏报代价远高于一般误报，仅用通用指标评价模型可能无法体现安全系统的业务价值。通过引入安全代价评价，可以更合理地解释专家规则融合模型的意义。

最后，从系统实现角度看，本文完成了从数据模拟、特征工程、模型训练、实验评估到 Web 可视化展示的完整流程，并增加智能解读、告警研判和用户画像模块。系统具有较强的演示性和可扩展性，可作为企业用户行为风险分析系统的原型。

## 1.3 国内外研究现状

用户行为分析是安全分析领域的重要方向。国外较早将 UBA 应用于内部威胁检测、账号盗用识别和异常访问审计中，常见方法包括统计基线、聚类分析、异常检测、序列建模和图分析等。统计基线方法通常根据用户历史行为计算访问时间、访问频率、资源类型等特征的正常范围，当新行为明显偏离基线时触发告警。该方法解释性较好，但对复杂模式识别能力有限。机器学习方法能够利用多维特征学习正常与异常之间的差异，常见算法包括逻辑回归、随机森林、梯度提升树、孤立森林、支持向量机以及深度学习模型等。这些方法提升了检测能力，但在落地中常面临样本标注不足、模型解释困难和误报控制等问题。

在内部威胁检测研究中，研究者通常关注用户登录行为、文件访问行为、邮件发送行为、设备使用行为和网络连接行为等多源数据。部分研究利用图模型描述用户、设备、资源之间的关系，通过关系异常发现潜在攻击链；也有研究将行为日志转化为时间序列，分析用户行为在时间维度上的变化。近年来，可解释机器学习逐渐受到重视，SHAP、LIME 等方法被用于解释模型预测结果，使安全人员能够理解某一条告警为何被判为高风险。

国内相关研究主要集中在日志审计、异常行为检测、数据泄露防护和态势感知等方向。随着等级保护、数据安全法和企业安全运营需求的发展，越来越多系统开始强调对账号、权限、数据访问和业务行为的综合分析。实际应用中，企业安全系统往往采用“规则 + 模型”的混合方式。一方面，规则用于表达明确的安全策略，例如深夜访问敏感数据、异地登录失败次数异常、越权操作等；另一方面，模型用于学习大量历史数据中的复杂统计规律。二者结合能够兼顾业务可控性和检测灵活性。

现有研究仍存在一些不足。第一，部分研究更关注算法指标，而对安全业务指标讨论不足，难以解释模型在实际运营中的价值。第二，部分系统输出结果停留在风险分数层面，缺少面向安全人员的自然语言解释和处置建议。第三，部分可视化系统仅展示静态报表，缺少从实时日志、用户聚合、关系图到实验对比的完整展示链路。本文针对这些不足，设计了融合 LightGBM、专家规则、SHAP 和智能文本研判的 UBA 系统。

## 1.4 本文研究内容

本文围绕企业用户行为风险分析系统展开，主要研究内容包括以下几个方面。

第一，构建用户行为日志模拟与特征工程流程。系统模拟企业中不同角色用户的行为日志，包含普通员工、财务人员、技术人员和管理员等角色，并注入远程异常登录、深夜高敏操作、批量下载、角色越权和操作频率异常等风险场景。在此基础上提取时间特征、频次特征、敏感操作特征、登录失败特征、用户基线偏离特征和角色操作匹配特征。

第二，设计基于 LightGBM 的异常检测模型。LightGBM 适合处理结构化表格特征，具有训练速度快、非线性表达能力强和特征重要性可分析等优点。系统采用时间切分方式构建训练集和测试集，避免随机切分带来的时间泄露问题。

第三，设计专家规则融合机制。针对安全业务中更关注的高危场景，系统设计四类专家规则，对模型风险分数进行上调。规则并不追求让所有通用指标最优，而是用于降低高危漏报、增强策略一致性和补充业务安全知识。

第四，设计模型解释与智能研判机制。系统通过 SHAP 输出单条日志的主要风险因素，并结合规则命中结果生成告警解释。进一步地，系统支持接入 DeepSeek / OpenAI 兼容接口，将结构化检测结果转化为自然语言研判文本，包括风险场景、置信度、关键证据和处置建议。

第五，实现可视化 Web 系统。系统后端基于 Flask 提供状态、实时日志、用户列表、用户详情、风险关系图、实验结果和特征重要性等接口；前端基于 ECharts 和 vis-network 展示实时风险仪表盘、行为雷达图、SHAP 柱状图、异常用户表、关系网络图和实验对比页面。

第六，开展实验评估与业务指标分析。本文对逻辑回归、孤立森林、LightGBM 和 LightGBM + Expert Rules 进行对比，分析通用指标和安全业务指标之间的关系，并说明专家规则融合模型在高危场景中的业务价值。

## 1.5 论文组织结构

本文共分为七章。第一章为绪论，介绍研究背景、研究意义、国内外研究现状和本文主要工作。第二章对系统进行需求分析，包括功能需求、非功能需求和数据安全需求。第三章介绍系统涉及的关键技术，包括 UBA、LightGBM、专家规则、SHAP、Flask、ECharts、vis-network 和大模型接口。第四章阐述系统总体设计，包括架构设计、数据流程、模型流程和数据库/文件组织。第五章介绍系统实现，包括数据模拟、特征工程、模型训练、后端接口、前端可视化和智能研判模块。第六章进行系统测试与实验分析，包括实验数据、评价指标、模型对比、混淆矩阵、安全业务指标和可解释性分析。第七章总结全文并展望后续改进方向。

# 2 系统需求分析

## 2.1 系统目标

本文系统的总体目标是构建一个能够对企业用户行为日志进行风险评估、异常检测、结果解释和可视化展示的安全防护原型系统。系统面向企业安全管理人员和数据安全审计人员，帮助其发现潜在异常行为、定位高风险用户、理解告警原因并形成初步处置建议。

从业务目标看，系统应能够识别用户在敏感数据访问、异常登录、越权操作和高频操作等方面的风险。系统并不以替代安全人员为目标，而是作为风险发现和辅助研判工具，为安全人员提供更快的线索筛选能力。从技术目标看，系统应完成从日志数据到模型预测再到前端展示的闭环流程，保证每个模块之间数据格式清晰、接口稳定、结果可复现。

## 2.2 功能需求分析

系统功能需求主要包括日志数据管理、风险检测、模型解释、用户聚合分析、关系图展示、实验对比和智能研判等模块。

日志数据管理模块负责读取和处理用户行为日志。系统当前采用模拟数据方式生成日志，但字段设计尽量贴近企业实际场景，包括用户编号、角色、时间、IP 位置、操作类型、文件敏感等级、登录失败情况等。模拟数据用于支撑模型训练、实验评估和前端演示。

风险检测模块负责对单条日志进行风险评分。系统先将原始日志转化为模型特征，再由 LightGBM 输出基础风险概率，随后根据专家规则对特定高危场景进行风险上调，最终得到综合风险分数。当前系统以 0.5 作为告警阈值，风险分数大于等于阈值时标记为告警。

模型解释模块负责说明风险判断原因。系统利用 SHAP 计算特征贡献，输出对风险分数影响最大的若干特征，例如角色操作不匹配、远程或未知位置、操作类型、文件敏感级别和登录失败次数等。同时，系统输出专家规则命中情况，使用户能够看到风险上调是否来自明确业务规则。

用户聚合分析模块负责从用户维度汇总风险。单条日志只能反映一次行为，而安全运营更关心一个用户在一段时间内的整体风险状态。系统对每个用户最近多条日志进行推理，计算平均风险、最高风险、告警次数、真实异常次数和代表性 SHAP 原因，并在异常检测表和用户详情页展示。

风险关系图模块用于展示用户与操作类型之间的关系。系统将用户节点和操作类型节点构成二分图，通过节点颜色、大小和边宽展示风险状态与操作强度。该模块便于安全人员从关系角度观察高风险用户集中关联的操作类别。

实验对比模块用于展示模型效果。系统提供通用机器学习指标、ROC/PR 曲线、混淆矩阵、安全业务指标和特征重要性排行。该模块不仅服务于系统开发过程，也服务于论文实验分析和答辩展示。

智能研判模块负责将结构化检测结果转化为自然语言说明。系统支持本地模板解释，也支持通过 `.env` 配置 DeepSeek / OpenAI 兼容接口。该模块输出智能解读、告警场景、关键证据、处置建议和用户画像文本，提升告警可读性。

## 2.3 非功能需求分析

系统的非功能需求主要包括准确性、实时性、可解释性、稳定性、可扩展性和安全性。

准确性方面，系统需要尽可能识别异常行为并减少误报。由于异常检测场景中正负样本不平衡，仅看准确率容易产生偏差。因此系统应同时关注 Precision、Recall、F1、AUC、PR-AUC 等指标，并结合安全业务指标评价高危场景漏报情况。

实时性方面，系统需要支持前端周期性轮询日志并快速返回风险分析结果。当前系统在启动时预加载模型、特征数据和 SHAP 解释器，并预计算用户聚合风险，以换取运行阶段更稳定的接口响应。虽然启动时间相对较长，但演示和使用阶段能够较快返回结果。

可解释性方面，系统不能只输出风险分数，还应解释风险来源。安全人员需要根据解释结果判断告警是否可信、是否需要进一步调查。系统通过 SHAP 和专家规则共同提供解释，使统计模型和业务规则能够互相补充。

稳定性方面，系统应保证数据文件、模型文件和前端依赖存在时能够正常启动。为提高答辩和演示稳定性，系统增加了启动检查模块，对特征数据、模型文件、实验结果、本地前端依赖和大模型配置进行提示。

可扩展性方面，系统后续可以接入真实日志数据，也可以增加新的规则、模型和前端页面。当前项目采用模块化脚本组织，`simulation.py`、`feature_engineering.py`、`model_training.py`、`evaluation_compare.py` 和 `app.py` 分别承担不同职责，便于后续维护。

安全性方面，系统需要避免泄露大模型 API Key 等敏感配置。项目使用 `.env` 管理 DeepSeek 接入参数，并在文档和答辩展示中避免暴露真实 Key。实时大模型调用也通过 `LLM_REALTIME_EVERY_N` 控制频率，避免持续轮询造成不必要的 token 消耗。同时，系统需求中需要重视“安全间隙”的识别：当第三方系统、远程访问、跨部门账号和敏感数据资源共同存在时，权限配置正确并不意味着行为一定安全。系统应通过持续日志分析识别权限使用过程中的异常偏离，并将检测结果转化为可复核的告警证据。

## 2.4 用户角色分析

系统主要面向三类用户。第一类是企业安全管理员，负责查看实时告警、识别高风险用户、分析风险原因并进行处置。该类用户需要直观的可视化界面、清晰的告警解释和可操作的处置建议。第二类是数据安全审计人员，负责从长期视角分析用户行为和权限使用情况。该类用户更关注用户画像、历史行为风险、混淆矩阵和模型评估结果。第三类是系统维护人员，负责数据更新、模型重训、接口维护和前端依赖管理。该类用户需要清晰的运行命令、启动检查和模块化代码结构。

## 2.5 数据需求分析

用户行为日志应包含能够描述用户身份、时间、位置、操作和资源敏感程度的信息。本文系统使用的数据规模为 {support['feature_rows']} 条特征样本，包含 {support['user_count']} 个用户，提取模型特征 {support['feature_count']} 个。整体异常样本占比约为 {pct(support['positive_ratio'])}，在测试集中异常样本数为 {support['experiment']['positive_count']}，正常样本数为 {support['experiment']['negative_count']}。这种样本分布符合安全异常检测中异常样本相对较少的特点。

为了支持模型训练和解释，数据字段需要经过数值化和特征工程处理。角色、IP 位置、操作类型等类别变量被编码为模型可接受的形式；时间戳被转化为小时、星期、是否周末、是否深夜等时间特征；窗口统计特征用于描述近一小时操作频次、敏感操作次数、登录失败次数和操作复杂度；用户基线偏离特征用于描述当前行为相对于用户历史平均水平的偏移。

# 3 关键技术介绍

## 3.1 用户行为分析技术

用户行为分析是一类以用户为中心的安全分析技术。其核心思想是通过收集用户在信息系统中的行为数据，建立用户正常行为基线，并在新行为出现时判断其是否偏离基线或命中高风险模式。异常检测综述研究指出，异常行为识别通常需要结合统计偏离、上下文特征和具体业务场景进行分析[2]。与传统基于攻击特征的检测方法相比，UBA 更关注“合法身份下的不正常行为”。在企业内部威胁、账号盗用和数据泄露防护场景中，UBA 具有较强适用性[4-5]。

UBA 的基本流程包括数据采集、数据清洗、特征提取、模型训练、风险评分、结果解释和告警展示。数据采集阶段需要收集登录日志、文件访问日志、操作日志和权限变更日志等；特征提取阶段需要将原始日志转化为可计算指标；模型训练阶段根据历史数据建立异常检测模型；风险评分阶段对新日志进行预测；结果解释阶段说明风险来源；告警展示阶段向安全人员提供可视化界面。

本文系统采用模拟日志作为数据来源，但保留了 UBA 系统的关键流程。系统不仅分析单条日志，还在用户维度进行聚合，使安全人员能够同时看到单次行为风险和长期用户风险。

## 3.2 LightGBM 算法

LightGBM 是一种基于梯度提升决策树的高效机器学习框架，适合处理结构化表格数据。其核心思想是通过多棵决策树逐步拟合损失函数的负梯度，从而得到具有较强非线性表达能力的集成模型[7-8]。相比传统 GBDT 实现，LightGBM 在训练速度、内存占用和大规模特征处理方面具有优势。

在用户行为风险检测中，特征往往包含类别编码、时间统计、窗口计数和连续数值等多种形式。LightGBM 能够较好处理这些混合特征，并捕捉特征之间的非线性组合关系。例如，深夜访问行为本身未必一定异常，但当深夜访问与高敏文件、远程位置和登录失败叠加出现时，风险显著升高。树模型能够通过分裂条件组合表达这种关系。

本文系统使用 LightGBM 作为主要监督学习模型。训练过程中考虑异常样本比例较低的问题，通过类别权重参数增强模型对异常样本的关注。模型训练完成后保存为 `lgbm_uba_model.pkl`，并将特征名保存为 `feature_names.pkl`，供后端服务加载和推理。

## 3.3 专家规则融合机制

专家规则是安全系统中常见的知识表达方式。相比机器学习模型，规则具有语义清晰、可控性强和便于审计的特点。本文并未将专家规则作为独立检测器，而是将其作为模型风险分数的补充机制。系统先由 LightGBM 输出基础概率，再根据规则命中情况增加风险分数，最后将分数限制在 0 到 1 之间。

当前系统设置了四类专家规则。第一类是深夜高敏操作，当 `is_night >= 0.5` 且 `file_sensitive_level >= 0.6` 时，将风险上调 0.4。第二类是周末高频操作，当 `is_weekend >= 0.5` 且 `op_count_1h >= 0.8` 时，将风险上调 0.3。第三类是远程或未知位置叠加登录失败，当 `is_remote_or_unknown_ip >= 0.5` 且 `login_fail_count_1h >= 0.4` 时，将风险上调 0.25。第四类是角色操作不匹配叠加高敏资源访问，当 `role_operation_mismatch >= 0.5` 且 `file_sensitive_level >= 0.6` 时，将风险上调 0.25。

专家规则融合的意义在于体现安全业务目标。安全系统并不总是追求平均指标最优，而是更关注关键风险是否被漏掉。通过规则兜底，系统能够对安全人员明确关心的场景保持敏感，从而降低高危漏报。

## 3.4 SHAP 可解释性方法

SHAP 是一种基于 Shapley Value 的模型解释方法。它将模型预测结果分解为各个特征的贡献值，用于说明某个特征对最终预测结果产生了正向还是负向影响[9-10]。在安全检测系统中，SHAP 可以帮助安全人员理解模型判断依据。

本文系统在后端加载模型时初始化 SHAP 解释器，对实时日志和用户代表性日志计算主要贡献特征。对于单条日志，系统返回贡献值绝对值最大的若干特征，并在前端以柱状图形式展示。对于用户聚合风险，系统选取最高风险日志作为代表样本，展示该用户主要风险因素。通过这种方式，系统可以回答“模型为什么认为该行为风险较高”这一问题。

需要说明的是，SHAP 主要解释 LightGBM 模型部分，而专家规则用于解释业务规则部分。两者共同构成双层解释机制：模型解释负责说明统计学习依据，规则解释负责说明安全策略依据。

## 3.5 Flask 与前后端交互

Flask 是 Python 生态中轻量级 Web 框架，适合快速构建 API 服务。本文系统后端使用 Flask 提供静态页面托管和 JSON 接口，包括 `/api/status`、`/api/next_log`、`/api/users`、`/api/user/<user_id>`、`/api/graph`、`/api/experiment_results` 和 `/api/feature_importance` 等。

前端通过周期性请求接口获取最新日志分析结果，并利用 ECharts 绘制仪表盘、雷达图、柱状图、模型对比图和曲线图。风险关系图使用 vis-network 实现交互式网络布局。前后端采用相对简单的接口方式，有利于降低系统复杂度，也便于答辩演示和后续扩展。

## 3.6 大模型智能解读

本文系统支持通过 DeepSeek / OpenAI 兼容接口增强自然语言解释能力。大模型模块不是检测模型的核心，而是用于将结构化检测结果转化为更易读的研判文本。输入信息包括模型基础分数、规则命中、SHAP 贡献、用户编号、时间戳和原始特征等；输出包括智能解读、风险场景、置信度、关键证据、处置建议和用户画像总结。

考虑到实时监控页面会持续轮询，如果每条日志都调用大模型，会造成较高 token 消耗并增加网络不稳定风险。因此系统提供 `LLM_REALTIME_EVERY_N` 参数，默认每 3 条实时日志调用一次大模型。若将 `LLM_DISABLED` 设置为 1，系统会回退到本地模板解释。这种设计兼顾了智能化效果和演示稳定性。

# 4 系统总体设计

## 4.1 系统总体架构

系统总体架构由数据层、算法层、服务层、展示层和智能增强层构成。数据层负责日志生成和特征数据管理；算法层负责模型训练、专家规则融合和 SHAP 解释；服务层负责 Flask API；展示层负责 Web 看板；智能增强层负责自然语言解读、告警研判和用户画像。

{md_image('fig1_system_architecture.png', '图4.1 系统总体架构')}

从图4.1可以看出，系统的核心不是单一模型，而是围绕用户行为风险分析构建的一套闭环流程。数据首先进入特征工程模块，形成训练和推理所需的特征表；模型训练模块输出 LightGBM 模型和特征名称；后端服务加载模型、特征数据和解释器，对实时样本进行推理；前端页面调用 API 展示风险结果；智能增强模块在结构化结果基础上生成自然语言解释。

## 4.2 系统运行流程

系统运行流程包括离线训练流程和在线展示流程。离线训练流程从 `simulation.py` 开始，生成用户行为日志；随后 `feature_engineering.py` 提取模型特征；`model_training.py` 训练 LightGBM 模型并保存产物；`evaluation_compare.py` 进行多模型对比实验并生成实验结果。在线展示流程由 `app.py` 启动 Flask 服务，加载模型和数据，向前端提供接口。

{md_image('fig2_pipeline.png', '图4.2 系统运行流水线')}

为了方便复现，项目提供了 `run_pipeline.py` 作为流水线入口。默认情况下，该脚本会依次执行数据模拟、特征工程、模型训练和实验评估。如果已有日志数据，也可以使用 `--skip-simulation` 参数跳过数据模拟，仅重建特征、模型和实验结果。

## 4.3 数据流程设计

系统数据流程从原始日志开始。每条日志包含用户编号、时间、角色、IP 位置、操作类型、文件敏感级别和异常标签等字段。特征工程模块将原始日志转化为模型特征，主要包括五类：基础类别特征、时间特征、窗口统计特征、用户基线偏离特征和业务规则相关特征。

基础类别特征包括用户角色编码、IP 位置编码和操作类型编码。时间特征包括小时、星期、是否周末、是否深夜和是否非工作时间。窗口统计特征包括近一小时操作次数、敏感操作次数、登录失败次数、敏感下载次数、不同操作类型数量和操作熵。用户基线偏离特征用于衡量当前操作频率和文件敏感级别相对于用户历史行为的偏移。业务规则相关特征包括远程或未知位置标记、角色操作不匹配标记等。

特征表以 parquet 格式保存，能够提高读取效率并保持字段类型稳定。模型训练产物保存在 `models/` 目录，实验结果保存在 `data/experiment_results.json`，前端展示所需图表和接口均基于这些产物生成。

## 4.4 模型融合设计

系统采用“模型基础分数 + 专家规则调整”的融合方式。设 LightGBM 输出的异常概率为 `p`，专家规则命中的风险调整量分别为 `r1, r2, ..., rn`，则综合风险分数可表示为：

```text
risk = min(1.0, p + sum(ri))
```

这种融合方式实现简单、解释直观，适合作为毕业设计系统的核心方案。其优点在于保留了模型从数据中学习复杂模式的能力，同时允许安全策略对特定场景进行显式约束。缺点是规则权重需要根据业务经验设定，后续可以通过验证集或安全运营反馈进一步优化。

## 4.5 前端页面设计

前端页面主要包含六个标签页。实时监控页展示综合风险评分、行为基线雷达图、SHAP 可解释性图、历史日志快照和智能解读面板。异常检测表展示所有用户的风险排序，支持按平均风险、最高风险、告警次数和日志数量排序。用户详情页展示指定用户的风险摘要、AI 用户画像、特征数据、SHAP 解释和最近日志记录。风险关系图展示用户与操作类别之间的关联。实验对比页展示模型性能、告警质量、ROC/PR 曲线、混淆矩阵、安全业务指标和特征重要性。系统原理页用于说明系统整体流程和关键技术。

这种页面组织方式既满足安全人员日常查看告警的需求，也满足论文答辩中展示系统完整性的需求。实时监控页强调动态演示，实验对比页强调模型评估，用户详情页强调个体解释，风险关系图强调关联分析。

# 5 系统详细实现

## 5.1 数据模拟实现

由于真实企业安全日志涉及隐私和敏感信息，本文采用模拟数据构建实验环境。数据模拟模块根据不同角色用户生成正常行为，并注入异常行为场景。正常行为遵循一定业务规律，例如普通员工主要进行常规文件查看和下载，财务人员可能访问财务相关资源，技术人员可能执行系统维护类操作，管理员具有更高权限。异常行为则模拟安全事件中常见的偏离模式。

异常场景主要包括异地或未知位置登录、深夜访问高敏文件、短时间批量下载、角色越权操作、登录失败聚集和操作复杂度异常等。通过这些场景，系统能够训练模型识别不同类型的风险行为。模拟数据虽然不能完全替代真实数据，但可以保证实验可控性和系统演示完整性。

## 5.2 特征工程实现

特征工程是 UBA 系统的关键环节。原始日志字段往往不能直接用于模型训练，需要转化为能够表达行为风险的数值特征。本文系统从时间、频次、敏感度、位置、角色匹配和用户基线等维度构造特征。

时间维度特征用于捕捉非工作时间行为。深夜、周末和下班后操作并不一定异常，但当这些时间特征与敏感文件访问、远程位置和高频操作结合时，风险会显著提高。频次维度特征用于描述短时间内的行为密度，例如近一小时操作次数、敏感操作次数和敏感下载次数。位置维度特征用于判断访问来源是否为远程或未知位置。角色匹配特征用于判断用户当前操作是否符合角色权限预期。用户基线偏离特征用于衡量当前行为与用户历史行为是否存在明显偏移。

通过这些特征，系统能够将原始日志中的离散事件转化为模型可学习的行为画像。最终生成的 `train_features.parquet` 文件既用于模型训练，也用于后端模拟实时日志流。

## 5.3 模型训练实现

模型训练模块读取特征数据，按照时间顺序划分训练集和测试集。时间切分更接近真实业务场景，因为安全系统通常使用历史数据训练模型，再对未来行为进行检测。如果随机切分，训练集和测试集可能混合相近时间段的行为，从而高估模型效果。

训练过程中，系统将 `role_code`、`ip_location_code` 和 `operation_type_code` 作为类别特征处理。由于异常样本比例较低，系统计算正负样本比例并设置类别权重，使模型在训练时更关注异常样本。模型训练完成后，系统保存模型文件、特征名称和训练元数据。训练元数据记录了时间切分、测试比例、类别特征、类别权重、阈值和专家规则参数，便于复现实验。

## 5.4 后端接口实现

后端 `app.py` 启动时会读取 `.env` 配置，加载 LightGBM 模型、特征名称、特征数据和 SHAP 解释器。随后系统预计算用户聚合风险数据，以便用户列表和用户详情接口能够快速返回结果。后端主要接口包括系统状态接口、实时日志接口、用户列表接口、用户详情接口、风险关系图接口、模型对比接口、特征重要性接口和实验结果接口。

实时日志接口 `/api/next_log` 是前端监控页的核心。该接口从特征数据中随机抽取样本，优先保证一定比例异常样本用于演示；随后计算基础模型分数、专家规则调整分数、规则命中、SHAP 贡献、告警研判和智能解读；最后以 JSON 格式返回前端。用户详情接口则基于预计算数据返回用户维度的平均风险、最高风险、告警次数、主要风险因素和最近日志。

为了提高系统演示稳定性，后端增加启动检查模块。启动时会检查特征数据、模型文件、特征列表、实验结果、本地 ECharts 和 vis-network 文件以及 LLM 配置状态。若文件缺失，系统会给出清晰提示，便于快速定位问题。

## 5.5 前端可视化实现

前端页面由 `web/index.html`、`web/static/app.js` 和 `web/static/app.css` 构成。页面采用标签页组织不同功能模块。实时监控页每隔一定时间调用后端接口，刷新风险仪表盘、行为雷达图、SHAP 柱状图和智能解读文本。异常检测表通过用户列表接口展示所有用户风险排序，并支持搜索、排序和阈值筛选。用户详情页通过用户详情接口展示单个用户的风险画像。风险关系图通过 vis-network 展示用户和操作类型之间的关联。

实验对比页是本轮系统增强的重要部分。页面不仅展示模型指标表，还展示实验结论摘要、中文指标解释、综合性能对比图、告警质量对比图、ROC/PR 曲线说明、风险因素排行、安全业务指标和混淆矩阵。相比单纯展示数值，这种设计更有利于论文说明和答辩表达。

## 5.6 智能解读与告警研判实现

智能解读模块接收模型分数、规则命中、SHAP 贡献和原始特征，将其转换为自然语言说明。在本地模板模式下，系统根据风险等级、规则命中和主要贡献特征拼接解释文本；在大模型模式下，系统通过 OpenAI 兼容接口请求 DeepSeek 等模型生成更自然的研判内容。

告警研判模块进一步将风险归因为具体安全场景，例如疑似账号盗用或暴力破解、疑似深夜高敏数据访问、疑似越权访问、疑似横向探索或扫描式操作、远程或未知位置访问异常等。研判结果包含置信度、关键证据和处置建议，使系统从风险打分工具扩展为安全分析辅助决策工具。

用户画像模块则从用户维度总结风险。它根据用户平均风险、最高风险、告警次数、最近日志和主要 SHAP 因素生成画像文本，帮助安全人员快速了解某个用户的整体行为状态。

# 6 系统测试与实验分析

## 6.1 实验环境与数据集

本文实验基于项目生成的用户行为特征数据进行。特征数据共 {support['feature_rows']} 条，用户数量为 {support['user_count']}，模型特征数量为 {support['feature_count']}。实验采用时间切分方式，测试集比例为 20%，测试样本数量为 {support['experiment']['test_count']}，其中异常样本 {support['experiment']['positive_count']} 条，正常样本 {support['experiment']['negative_count']} 条，异常占比为 {pct(support['experiment']['positive_ratio'])}。

测试集异常占比不高，符合安全检测任务中异常行为相对少见的特点。在这种情况下，Accuracy 可能不能充分反映模型对异常样本的识别能力，因此本文同时使用 Precision、Recall、F1、AUC 和 PR-AUC 等指标，并补充高危场景召回率和安全代价指标。

## 6.2 评价指标设计

通用评价指标包括 Accuracy、Precision、Recall、F1-Score、AUC 和 PR-AUC。Accuracy 表示整体判断正确比例，但当正常样本远多于异常样本时，准确率可能偏乐观。Precision 表示被系统判为异常的样本中真实异常的比例，反映误报压力。Recall 表示真实异常样本中被系统检出的比例，反映漏报风险。F1-Score 是 Precision 和 Recall 的调和平均，用于综合衡量告警质量。AUC 表示模型区分正常与异常的整体能力，PR-AUC 在异常样本占比较低时更能体现告警质量。

为了贴合数据安全业务，本文进一步设计安全业务指标。高危场景召回率 HighRisk-Recall 用于衡量模型对关键风险场景的识别能力。加权安全代价 Security-Cost 定义如下：

```text
Security-Cost = 5 * 高危漏报 + 3 * 普通漏报 + 1 * 误报
```

该公式体现了高危漏报代价高于普通漏报，普通漏报代价高于误报的安全业务假设。Security-Score 则基于安全代价归一化得到，分数越高表示业务安全收益越好。

## 6.3 多模型实验结果

本文对逻辑回归、孤立森林、LightGBM 和 LightGBM + Expert Rules 四种方法进行对比。逻辑回归作为线性分类基线，孤立森林作为无监督异常检测基线，LightGBM 作为主要监督学习模型，LightGBM + Expert Rules 作为融合专家规则的安全增强模型。

{md_image('fig3_model_metrics.png', '图6.1 模型通用性能指标对比')}

从图6.1和表6.1可以看出，LightGBM 在通用指标上表现最好，其 Accuracy 为 {lgb['Accuracy']:.4f}，Precision 为 {lgb['Precision']:.4f}，Recall 为 {lgb['Recall']:.4f}，F1-Score 为 {lgb['F1-Score']:.4f}，PR-AUC 为 {lgb['PR-AUC']:.4f}。逻辑回归也取得了较高指标，说明部分异常行为在线性特征空间中已经具有一定可分性。孤立森林表现相对较弱，说明在本文构造的监督标签场景中，无监督方法难以充分利用异常样本信息。

LightGBM + Expert Rules 的 F1-Score 为 {hybrid['F1-Score']:.4f}，略低于纯 LightGBM；Precision 也从 {lgb['Precision']:.4f} 降至 {hybrid['Precision']:.4f}。这说明专家规则上调风险后会带来一定误报增加。但从 Recall 看，融合模型由 {lgb['Recall']:.4f} 提升至 {hybrid['Recall']:.4f}，说明规则增强有助于减少漏报。

表6.1 模型实验结果对比

| 模型 | Accuracy | Precision | Recall | F1 | PR-AUC | HighRisk-Recall | Security-Cost |
|---|---:|---:|---:|---:|---:|---:|---:|
| Logistic Regression | {lr['Accuracy']:.4f} | {lr['Precision']:.4f} | {lr['Recall']:.4f} | {lr['F1-Score']:.4f} | {lr['PR-AUC']:.4f} | {lr['HighRisk-Recall']:.4f} | {lr['Security-Cost']:.0f} |
| Isolation Forest | {iso['Accuracy']:.4f} | {iso['Precision']:.4f} | {iso['Recall']:.4f} | {iso['F1-Score']:.4f} | {iso['PR-AUC']:.4f} | {iso['HighRisk-Recall']:.4f} | {iso['Security-Cost']:.0f} |
| LightGBM | {lgb['Accuracy']:.4f} | {lgb['Precision']:.4f} | {lgb['Recall']:.4f} | {lgb['F1-Score']:.4f} | {lgb['PR-AUC']:.4f} | {lgb['HighRisk-Recall']:.4f} | {lgb['Security-Cost']:.0f} |
| LightGBM + Expert Rules | {hybrid['Accuracy']:.4f} | {hybrid['Precision']:.4f} | {hybrid['Recall']:.4f} | {hybrid['F1-Score']:.4f} | {hybrid['PR-AUC']:.4f} | {hybrid['HighRisk-Recall']:.4f} | {hybrid['Security-Cost']:.0f} |

## 6.4 安全业务指标分析

在数据安全场景中，高危漏报往往比普通误报更严重。一次高危漏报可能意味着敏感数据被异常访问而未被发现，而一次误报通常只会增加安全人员复核成本。因此，仅从 F1 或 AUC 判断模型优劣是不充分的。

{md_image('fig4_security_metrics.png', '图6.2 高危召回率与安全代价对比')}

从图6.2可以看出，LightGBM + Expert Rules 的高危场景召回率达到 {hybrid['HighRisk-Recall']:.4f}，高于纯 LightGBM 的 {lgb['HighRisk-Recall']:.4f}；加权安全代价从 {lgb['Security-Cost']:.0f} 降低至 {hybrid['Security-Cost']:.0f}。这说明专家规则虽然会略微降低部分通用指标，但能够在安全业务更关心的高危场景中发挥兜底作用。

这一结果也说明本文系统对专家规则融合的定位是合理的。专家规则不是为了“刷高”所有指标，而是为了表达安全策略、降低高危漏报并增强解释能力。在实际企业安全运营中，这类能力往往比单一统计指标更有价值。

## 6.5 混淆矩阵分析

混淆矩阵能够直观展示模型误报和漏报情况。对于安全检测系统，右下角“异常判异常”代表成功检出，左下角“异常漏报正常”代表漏报，右上角“正常误报异常”代表误报。

{md_image('fig5_confusion_matrix.png', '图6.3 LightGBM + 专家规则混淆矩阵')}

LightGBM + Expert Rules 在测试集中将 {hybrid['confusion_matrix']['tp']} 条异常样本正确判为异常，将 {hybrid['confusion_matrix']['fn']} 条异常样本漏报为正常，同时产生 {hybrid['confusion_matrix']['fp']} 条误报。相比 LightGBM，该模型漏报数量由 {lgb['confusion_matrix']['fn']} 降至 {hybrid['confusion_matrix']['fn']}，但误报数量由 {lgb['confusion_matrix']['fp']} 增至 {hybrid['confusion_matrix']['fp']}。这体现了少漏报与少误报之间的取舍。

对于企业安全防护系统而言，在可控范围内增加少量误报以降低高危漏报是可以接受的。尤其是在敏感数据保护场景中，漏报带来的潜在损失往往远大于误报复核成本。

## 6.6 特征重要性分析

特征重要性能够反映模型主要依赖哪些行为因素进行判断。本文根据 LightGBM 的 Gain 指标绘制特征重要性排行。

{md_image('fig6_feature_importance.png', '图6.4 LightGBM 特征重要性 Top 10')}

从图6.4可以看出，角色操作不匹配、操作类型、远程或未知位置、文件敏感级别、敏感度偏离和登录失败次数等特征对模型判断具有较大影响。这与安全业务经验基本一致。角色操作不匹配通常意味着用户执行了不符合岗位职责的操作；远程或未知位置可能暗示账号被盗用或异常登录；文件敏感级别和敏感度偏离能够反映用户是否访问了超出日常范围的高价值数据；登录失败次数则与暴力破解或账号异常尝试有关。

特征重要性分析也说明本文特征工程设计较为有效。系统没有只依赖单一字段，而是从角色、位置、时间、操作、敏感资源和用户基线等多个角度综合判断风险。

## 6.7 可解释性与智能化效果分析

系统可解释性主要体现在三方面。第一，SHAP 输出模型特征贡献，使安全人员能够看到风险分数受到哪些特征影响。第二，专家规则输出明确命中原因，例如深夜高敏操作或远程登录失败聚集。第三，智能解读模块将结构化结果转化为自然语言说明，降低安全人员理解成本。

在实时监控页面中，系统会展示综合风险评分、行为雷达图和 SHAP 柱状图，并在左侧面板输出智能解读和告警研判结果。对于高风险日志，系统会给出处置建议，例如触发二次认证、临时收敛权限、保留会话审计记录或进入人工复核队列。对于低风险日志，系统则建议维持常规监控并纳入周期性基线比对。

用户画像模块进一步提升了用户维度的分析能力。安全人员可以查看某一用户的平均风险、最高风险、告警次数、真实异常次数、主要风险因素和最近日志记录。相比只看单条告警，用户画像更适合发现持续风险和行为趋势。

## 6.8 系统测试

系统测试主要包括功能测试、接口测试和运行稳定性测试。功能测试验证实时监控、异常用户表、用户详情、风险关系图和实验对比页面是否能够正常加载。接口测试验证 `/api/status`、`/api/next_log`、`/api/users`、`/api/user/<user_id>`、`/api/graph`、`/api/experiment_results` 和 `/api/feature_importance` 是否能够返回正确 JSON 数据。运行稳定性测试验证系统在加载模型、读取特征数据、初始化 SHAP 和预计算用户聚合数据后能否正常启动。

交接验证结果显示，核心接口均能够返回 200 状态码，说明后端服务链路基本稳定。系统启动较慢主要是因为需要加载模型、读取 100000 条特征数据、初始化 SHAP 并预计算 200 个用户的聚合风险，该现象属于当前实现的正常现象。为保证答辩稳定，建议演示前提前启动服务。

# 7 总结与展望

## 7.1 总结

本文围绕企业信息数据安全防护需求，设计并实现了一套基于用户行为数据评估的 UBA 系统。系统从用户行为日志出发，通过特征工程提取时间、位置、角色、操作、敏感资源和用户基线偏离等多维特征；采用 LightGBM 模型进行异常检测；引入专家规则对高危安全场景进行风险上调；利用 SHAP 方法解释模型判断依据；通过 Flask 提供后端接口，并基于 ECharts 和 vis-network 实现多维可视化展示；进一步接入 DeepSeek / OpenAI 兼容接口，实现智能解读、告警研判和用户画像总结。

实验结果表明，LightGBM 在通用指标上表现较好，F1-Score 达到 {lgb['F1-Score']:.4f}，PR-AUC 达到 {lgb['PR-AUC']:.4f}。专家规则融合模型虽然在 Precision 和 PR-AUC 上略有下降，但高危场景召回率提升至 {hybrid['HighRisk-Recall']:.4f}，加权安全代价降低至 {hybrid['Security-Cost']:.0f}。这说明在安全业务场景中，模型评价不能只看通用统计指标，还应关注高危漏报和业务代价。专家规则的价值在于安全兜底、策略一致性和解释增强。

本文系统的主要成果包括：构建了完整的数据模拟与特征工程流程；完成了 LightGBM 与专家规则融合的风险检测机制；设计了 SHAP 与规则结合的双层解释机制；实现了实时监控、异常用户表、用户画像、风险关系图和实验对比页面；设计了安全业务评价指标；实现了大模型辅助自然语言研判。整体来看，系统形成了“检测、解释、研判、处置建议”的闭环，具有一定实用价值和演示价值。

## 7.2 不足与展望

本文系统仍存在一些不足。首先，当前数据主要来自模拟生成，虽然字段和异常场景尽量贴近企业业务，但与真实生产环境仍有差距。后续可接入真实脱敏日志，进一步验证模型在真实数据上的泛化能力。其次，专家规则权重主要基于经验设定，后续可利用验证集、网格搜索或安全运营反馈自动调整规则权重。再次，当前用户聚合风险在启动时预计算，导致 Flask 启动时间较长。后续可引入缓存文件或异步计算机制，减少启动等待时间。

此外，当前大模型模块主要用于解释文本生成，尚未参与策略优化和交互式问答。后续可探索将大模型用于告警摘要、调查报告生成、规则推荐和处置流程编排。但在安全系统中，大模型输出必须受到约束，不能替代模型检测和人工决策。最后，前端 ECharts 依赖仍建议完全本地化，以保证答辩或离线环境下图表稳定加载。

综上，本文系统完成了企业用户行为风险分析的核心流程，并在模型解释、实验展示和智能研判方面进行了增强。后续若结合真实日志、权限系统和安全运营流程，可进一步发展为更完整的数据安全监测与告警研判平台。

# 参考文献

[1] ALMOHAIMEED M, ALBALWY F, ALHARBI R, et al. Developing a comprehensive cyber risk assessment framework for supply chains: insights into third-party vulnerabilities and security gaps[J]. Intelligent Information Management, 2025, 17: 58-77.

[2] CHANDOLA V, BANERJEE A, KUMAR V. Anomaly detection: a survey[J]. ACM Computing Surveys, 2009, 41(3): 1-58.

[3] BUCZAK A L, GUVEN E. A survey of data mining and machine learning methods for cyber security intrusion detection[J]. IEEE Communications Surveys & Tutorials, 2016, 18(2): 1153-1176.

[4] SALEM M B, HERSHKOP S, STOLFO S J. A survey of insider attack detection research[M]//Insider Attack and Cyber Security. Boston: Springer, 2008: 69-90.

[5] SHASHANKA M, SHEN M Y, WANG J. User and entity behavior analytics for enterprise security[C]//2016 IEEE International Conference on Big Data. Washington: IEEE, 2016: 1867-1874.

[6] LIU F T, TING K M, ZHOU Z H. Isolation forest[C]//2008 IEEE International Conference on Data Mining. Pisa: IEEE, 2008: 413-422.

[7] KE G, MENG Q, FINLEY T, et al. LightGBM: a highly efficient gradient boosting decision tree[C]//Advances in Neural Information Processing Systems 30. Long Beach: Curran Associates, 2017: 3146-3154.

[8] FRIEDMAN J H. Greedy function approximation: a gradient boosting machine[J]. The Annals of Statistics, 2001, 29(5): 1189-1232.

[9] LUNDBERG S M, LEE S I. A unified approach to interpreting model predictions[C]//Advances in Neural Information Processing Systems 30. Long Beach: Curran Associates, 2017: 4765-4774.

[10] LUNDBERG S M, ERION G, CHEN H, et al. From local explanations to global understanding with explainable AI for trees[J]. Nature Machine Intelligence, 2020, 2(1): 56-67.

[11] RIBEIRO M T, SINGH S, GUESTRIN C. Why should I trust you? Explaining the predictions of any classifier[C]//Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining. San Francisco: ACM, 2016: 1135-1144.

[12] SARKER I H, KAYES A S M, BADSHA S, et al. Cybersecurity data science: an overview from machine learning perspective[J]. Journal of Big Data, 2020, 7: 41.

[13] NIST. Guide for conducting risk assessments: NIST Special Publication 800-30 Revision 1[R]. Gaithersburg: National Institute of Standards and Technology, 2012.

[14] NIST. Security and privacy controls for information systems and organizations: NIST Special Publication 800-53 Revision 5[R]. Gaithersburg: National Institute of Standards and Technology, 2020.

[15] NIST. The NIST Cybersecurity Framework (CSF) 2.0[R]. Gaithersburg: National Institute of Standards and Technology, 2024.

[16] ISO/IEC. ISO/IEC 27001:2022 Information security, cybersecurity and privacy protection - Information security management systems - Requirements[S]. Geneva: International Organization for Standardization, 2022.

[17] GRINBERG M. Flask Web Development: Developing Web Applications with Python[M]. 2nd ed. Sebastopol: O'Reilly Media, 2018.

[18] MCKINNEY W. Data structures for statistical computing in Python[C]//Proceedings of the 9th Python in Science Conference. Austin: SciPy, 2010: 56-61.

[19] PEDREGOSA F, VAROQUAUX G, GRAMFORT A, et al. Scikit-learn: machine learning in Python[J]. Journal of Machine Learning Research, 2011, 12: 2825-2830.

[20] HUNTER J D. Matplotlib: a 2D graphics environment[J]. Computing in Science & Engineering, 2007, 9(3): 90-95.

# 致谢

在本论文和系统开发完成过程中，感谢指导教师在选题、系统设计、实验分析和论文撰写方面给予的指导。通过本次毕业设计，我对企业数据安全防护、用户行为分析、机器学习异常检测、可解释性方法和 Web 可视化开发有了更加系统的理解。

同时感谢开源社区提供的 Python、Flask、pandas、scikit-learn、LightGBM、SHAP、ECharts 和 vis-network 等工具。这些工具为系统实现提供了重要支持。最后感谢同学和朋友在系统测试、答辩演示准备和论文修改过程中提出的建议。由于本人能力有限，论文和系统仍存在不足，恳请各位老师批评指正。
"""
    return text


def md_to_docx(md: str) -> None:
    doc = Document()
    style_doc(doc)

    cover_lines = [
        ("吉 林 化 工 学 院", 22, True, 26),
        ("毕业设计（论文）", 22, True, 42),
        ("", 12, False, 10),
        ("设计（论文）题目：\t基于用户行为数据评估的", 15, False, 10),
        ("\t\t企业信息数据安全防护系统", 15, False, 34),
        ("学生姓名：\t__________\t学号：\t__________", 12, False, 12),
        ("学院：\t信息与控制工程学院", 12, False, 12),
        ("专业：\t计算机相关专业", 12, False, 12),
        ("班级：\t__________", 12, False, 12),
        ("校内指导教师：\t__________", 12, False, 12),
        ("校外指导教师：\t/\t/\t/", 12, False, 34),
        ("2026\t年\t05\t月", 12, False, 0),
    ]
    for text, size, bold, after in cover_lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_after = Pt(after)
        run = p.add_run(text)
        run.font.name = "黑体" if bold else "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体" if bold else "宋体")
        run.font.size = Pt(size)
        run.bold = bold
    doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目录")
    r.font.name = "黑体"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
    r.font.size = Pt(16)
    r.bold = True
    toc_items = [
        "摘要",
        "Abstract",
        "1 绪论",
        "2 系统需求分析",
        "3 关键技术介绍",
        "4 系统总体设计",
        "5 系统详细实现",
        "6 系统测试与实验分析",
        "7 总结与展望",
        "参考文献",
        "致谢",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(item)
        run.font.name = "宋体"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        run.font.size = Pt(12)
    doc.add_page_break()

    raw_lines = md.splitlines()
    start = next((i for i, line in enumerate(raw_lines) if line.strip() == "## 摘要"), 0)
    lines = raw_lines[start:]
    in_table = False
    table_lines: list[str] = []

    def flush_table():
        nonlocal table_lines
        if not table_lines:
            return
        rows = []
        for line in table_lines:
            if re.match(r"^\|\s*-", line):
                continue
            parts = [p.strip() for p in line.strip("|").split("|")]
            rows.append(parts)
        if rows:
            add_table(doc, rows[0], rows[1:])
        table_lines = []

    for line in lines:
        if line.startswith("|"):
            table_lines.append(line)
            continue
        flush_table()
        if line.startswith("# "):
            text = line[2:].strip()
            if text in {"1 绪论", "2 系统需求分析", "3 关键技术介绍", "4 系统总体设计", "5 系统详细实现", "6 系统测试与实验分析", "7 总结与展望", "参考文献", "致谢"}:
                doc.add_page_break()
            add_heading(doc, text, 1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), 3)
        elif line.startswith("!["):
            m = re.match(r"!\[(.*?)\]\((.*?)\)", line)
            if m:
                caption = m.group(1)
                path = Path(m.group(2))
                add_image(doc, path.name, caption)
        elif line.startswith("```"):
            continue
        elif not line.strip():
            continue
        else:
            if line.strip().startswith("[") and re.match(r"^\[\d+\]", line.strip()):
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Pt(0)
                p.paragraph_format.line_spacing = 1.5
                run = p.add_run(line.strip())
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
                run.font.size = Pt(10.5)
            else:
                add_para(doc, line.strip())
    flush_table()
    try:
        doc.save(OUT_DOCX)
    except PermissionError:
        doc.save(ALT_DOCX)


def main() -> None:
    support = load_support()
    md = build_markdown(support)
    OUT_MD.write_text(md, encoding="utf-8")
    md_to_docx(md)
    chinese_chars = len(re.findall(r"[\u4e00-\u9fff]", md))
    total_chars = len(re.sub(r"\s+", "", md))
    print(f"written: {OUT_MD}")
    print(f"written: {OUT_DOCX if OUT_DOCX.exists() else ALT_DOCX}")
    print(f"chinese_chars={chinese_chars}, non_space_chars={total_chars}")


if __name__ == "__main__":
    main()
