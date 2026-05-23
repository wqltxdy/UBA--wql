from __future__ import annotations

from pathlib import Path
from zipfile import ZipFile

from docx import Document


docx_path = Path("毕业论文初稿.docx")
md_path = Path("毕业论文初稿.md")
doc = Document(str(docx_path))
texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
with ZipFile(docx_path) as zf:
    media = [n for n in zf.namelist() if n.startswith("word/media/")]

print("docx_size", docx_path.stat().st_size)
print("md_size", md_path.stat().st_size)
print("paragraphs", len(doc.paragraphs))
print("nonempty_paragraphs", len(texts))
print("tables", len(doc.tables))
print("embedded_media", len(media))
print("first_items")
for item in texts[:25]:
    print(item)
