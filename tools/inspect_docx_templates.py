from __future__ import annotations

from docx import Document


for name in ["毕业论文参考样例1.docx", "毕业论文参考样例2.docx"]:
    doc = Document(name)
    print(f"\n==== {name} ====")
    print("paragraphs", len(doc.paragraphs), "tables", len(doc.tables), "sections", len(doc.sections))
    count = 0
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            print(i, para.style.name, repr(text[:120]))
            count += 1
        if count >= 90:
            break
    styles = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text and para.style.name not in styles:
            styles.append(para.style.name)
    print("styles", styles[:50])
    for si, section in enumerate(doc.sections):
        print(
            "section", si,
            "page", section.page_width, section.page_height,
            "margins", section.top_margin, section.bottom_margin, section.left_margin, section.right_margin,
            "header/footer", section.header_distance, section.footer_distance,
        )
